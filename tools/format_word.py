#!/usr/bin/env python3
"""
Post-process Word document formatting.
Applies professional formatting including page size, margins, and styles.
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION
except ImportError:
    print("❌ python-docx not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)

def format_word_document(doc_path):
    """
    Apply professional formatting to Word document.
    
    Args:
        doc_path: Path to Word document
    """
    if not os.path.exists(doc_path):
        print(f"❌ File not found: {doc_path}", file=sys.stderr)
        return 1
    
    try:
        doc = Document(doc_path)
        
        # Set page size to 6x9 inches (KDP Print standard)
        sections = doc.sections
        for section in sections:
            # Set page size to 6x9 inches
            section.page_height = Inches(9)
            section.page_width = Inches(6)
            
            # Set margins: Inside 0.375", Outside 0.25", Top/Bottom 0.25"
            section.top_margin = Inches(0.25)
            section.bottom_margin = Inches(0.25)
            section.left_margin = Inches(0.375)  # Inside (gutter)
            section.right_margin = Inches(0.25)  # Outside
            
            # Different margins for odd/even pages (mirror margins)
            section.different_first_page_header_footer = True
            section.odd_and_even_pages_header_footer = True
        
        # Apply professional formatting to paragraphs
        for para in doc.paragraphs:
            # Skip if empty
            if not para.text.strip():
                continue
            
            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                if para.style.name == 'Heading 1':
                    # Chapter titles - larger, bold, page break before
                    para.paragraph_format.page_break_before = True
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(12)
                    if para.runs:
                        para.runs[0].font.size = Pt(22)
                        para.runs[0].font.bold = True
                        para.runs[0].font.name = 'Arial'
                elif para.style.name == 'Heading 2':
                    para.paragraph_format.space_before = Pt(12)
                    para.paragraph_format.space_after = Pt(6)
                    if para.runs:
                        para.runs[0].font.size = Pt(15)
                        para.runs[0].font.bold = True
                        para.runs[0].font.name = 'Arial'
                elif para.style.name == 'Heading 3':
                    para.paragraph_format.space_before = Pt(10)
                    para.paragraph_format.space_after = Pt(5)
                    if para.runs:
                        para.runs[0].font.size = Pt(13)
                        para.runs[0].font.bold = True
                        para.runs[0].font.name = 'Arial'
            else:
                # Body text - professional formatting
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                para.paragraph_format.line_spacing = 1.45
                para.paragraph_format.first_line_indent = Inches(0.15)
                para.paragraph_format.space_after = Pt(3)
                
                # Set font to Georgia 10pt
                for run in para.runs:
                    run.font.name = 'Georgia'
                    run.font.size = Pt(10)
        
        # Format first paragraph after headings (no indent)
        prev_para = None
        for para in doc.paragraphs:
            if prev_para and prev_para.style.name.startswith('Heading'):
                # First paragraph after heading - no indent, slightly larger
                para.paragraph_format.first_line_indent = Inches(0)
                para.paragraph_format.space_after = Pt(6)
                for run in para.runs:
                    if run.font.size is None or run.font.size.pt < 10.5:
                        run.font.size = Pt(10.5)
            prev_para = para
        
        # Format blockquotes
        for para in doc.paragraphs:
            if 'blockquote' in para.style.name.lower() or para.style.name == 'Quote':
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.right_indent = Inches(0.3)
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(8)
                for run in para.runs:
                    run.font.italic = True
                    run.font.size = Pt(9.5)
                    run.font.name = 'Georgia'
        
        # Save the formatted document
        doc.save(doc_path)
        print(f"✅ Word document formatted: {doc_path}")
        return 0
        
    except Exception as e:
        print(f"❌ Error formatting Word document: {e}", file=sys.stderr)
        return 1

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: format_word.py <word_document.docx>", file=sys.stderr)
        sys.exit(1)
    
    doc_path = sys.argv[1]
    sys.exit(format_word_document(doc_path))
